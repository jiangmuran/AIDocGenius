import os
from pathlib import Path
from typing import Union, Any, Optional, Dict
import json
try:
    import yaml
    YAML_AVAILABLE = True
except ImportError:
    yaml = None
    YAML_AVAILABLE = False
from docx import Document
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PyPDF2 = None
    PDF_AVAILABLE = False
import logging

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def load_document(file_path: Union[str, Path]) -> Any:
    """
    加载文档内容，支持多种格式
    
    Args:
        file_path: 文档路径
        
    Returns:
        str: 文档内容
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()
    
    try:
        if suffix in ['.txt', '.rst']:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
                
        elif suffix == '.md':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
                
        elif suffix == '.docx':
            doc = Document(file_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
        elif suffix == '.pdf':
            if not PDF_AVAILABLE:
                raise ImportError("PyPDF2 is required to read PDF files")
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                return '\n'.join([page.extract_text() or "" for page in pdf_reader.pages])
                
        elif suffix in ['.json', '.yaml', '.yml']:
            with open(file_path, 'r', encoding='utf-8') as f:
                if suffix == '.json':
                    return json.load(f)
                else:
                    if not YAML_AVAILABLE:
                        raise ImportError("pyyaml is required to read YAML files")
                    return yaml.safe_load(f)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
            
    except Exception as e:
        logger.error(f"Error loading file {file_path}: {str(e)}")
        raise

def save_document(content: Any, file_path: Union[str, Path], format_options: dict = None) -> None:
    """
    保存文档内容
    
    Args:
        content: 要保存的内容
        file_path: 保存路径
        format_options: 格式选项
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()
    
    try:
        if suffix in ['.txt', '.rst']:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(str(content))
                
        elif suffix == '.md':
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
        
        elif suffix == '.html':
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(str(content))
                
        elif suffix == '.docx':
            if isinstance(content, Document):
                content.save(str(file_path))
            else:
                doc = Document()
                if isinstance(content, str):
                    # 按段落分割
                    paragraphs = content.split('\n\n')
                    for p in paragraphs:
                        if p.strip():
                            doc.add_paragraph(p.strip())
                else:
                    doc.add_paragraph(str(content))
                doc.save(str(file_path))
            
        elif suffix == '.pdf':
            # PDF generation requires additional libraries
            raise NotImplementedError("PDF generation not implemented yet")
            
        elif suffix in ['.json', '.yaml', '.yml']:
            with open(file_path, 'w', encoding='utf-8') as f:
                if suffix == '.json':
                    json.dump(content, f, ensure_ascii=False, indent=2)
                else:
                    if not YAML_AVAILABLE:
                        raise ImportError("pyyaml is required to write YAML files")
                    yaml.safe_dump(content, f, allow_unicode=True)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
            
    except Exception as e:
        logger.error(f"Error saving file {file_path}: {str(e)}")
        raise

def get_file_info(file_path: Union[str, Path]) -> dict:
    """
    获取文件信息
    
    Args:
        file_path: 文件路径
        
    Returns:
        dict: 文件信息
    """
    file_path = Path(file_path)
    
    return {
        'name': file_path.name,
        'extension': file_path.suffix,
        'size': os.path.getsize(file_path),
        'created_time': os.path.getctime(file_path),
        'modified_time': os.path.getmtime(file_path),
        'is_binary': is_binary_file(file_path)
    }

def is_binary_file(file_path: Union[str, Path]) -> bool:
    """
    检查文件是否为二进制文件
    
    Args:
        file_path: 文件路径
        
    Returns:
        bool: 是否为二进制文件
    """
    try:
        with open(file_path, 'tr') as f:
            f.read(1024)
            return False
    except:
        return True

def create_directory_if_not_exists(directory: Union[str, Path]) -> None:
    """
    如果目录不存在则创建
    
    Args:
        directory: 目录路径
    """
    Path(directory).mkdir(parents=True, exist_ok=True)

def ensure_text(content: Any) -> str:
    """
    将内容标准化为字符串，供摘要/分析/翻译使用
    """
    if isinstance(content, str):
        return content
    if isinstance(content, (dict, list)):
        return json.dumps(content, ensure_ascii=False, indent=2)
    return str(content)

def load_config(config_path: Optional[Union[str, Path]]) -> Dict[str, Any]:
    """
    加载配置文件（JSON/YAML）
    """
    if not config_path:
        return {}

    path = Path(config_path)
    if not path.exists():
        raise FileNotFoundError(f"Config file not found: {path}")

    suffix = path.suffix.lower()
    if suffix == '.json':
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)

    if suffix in ['.yaml', '.yml']:
        if not YAML_AVAILABLE:
            raise ImportError("pyyaml is required to read YAML config files")
        with open(path, 'r', encoding='utf-8') as f:
            return yaml.safe_load(f) or {}

    raise ValueError(f"Unsupported config format: {suffix}")
