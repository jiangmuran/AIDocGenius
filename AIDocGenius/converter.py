from pathlib import Path
from typing import Optional, Dict, Any
import markdown
from docx import Document
from .utils import load_document, save_document, logger

class Converter:
    """
    文档格式转换器
    """
    
    def __init__(self):
        """
        初始化转换器
        """
        self.supported_formats = {
            '.txt': self._convert_to_txt,
            '.md': self._convert_to_markdown,
            '.docx': self._convert_to_docx,
            '.html': self._convert_to_html,
            '.json': self._convert_to_json,
            '.yaml': self._convert_to_yaml
        }
        logger.info("Initialized converter")
        
    def convert(self,
               input_path: str,
               output_path: str,
               format_options: Optional[Dict[str, Any]] = None) -> None:
        """
        转换文档格式
        
        Args:
            input_path: 输入文件路径
            output_path: 输出文件路径
            format_options: 格式选项
        """
        input_path = Path(input_path)
        output_path = Path(output_path)
        
        try:
            # 检查格式支持
            if output_path.suffix.lower() not in self.supported_formats:
                raise ValueError(f"Unsupported output format: {output_path.suffix}")
                
            # 加载文档
            content = load_document(input_path)
            
            # 转换格式
            converter = self.supported_formats[output_path.suffix.lower()]
            converted_content = converter(content, format_options or {})
            
            # 保存文档
            save_document(converted_content, output_path, format_options)
            logger.info(f"Converted {input_path} to {output_path}")
            
        except Exception as e:
            logger.error(f"Conversion error: {str(e)}")
            raise
            
    def _convert_to_txt(self, content: Any, options: dict) -> str:
        """
        转换为纯文本格式
        """
        return str(content)
        
    def _convert_to_markdown(self, content: Any, options: dict) -> str:
        """
        转换为Markdown格式
        """
        if isinstance(content, str):
            return content
        elif isinstance(content, dict):
            # 将字典转换为Markdown表格
            md = "| Key | Value |\n|------|--------|\n"
            for key, value in content.items():
                md += f"| {key} | {value} |\n"
            return md
        else:
            return str(content)
            
    def _convert_to_docx(self, content: Any, options: dict) -> Document:
        """
        转换为Word文档格式
        """
        doc = Document()
        
        if isinstance(content, str):
            # 按段落分割
            paragraphs = content.split('\n\n')
            for p in paragraphs:
                doc.add_paragraph(p.strip())
        elif isinstance(content, dict):
            # 添加标题
            doc.add_heading('Document Content', 0)
            # 添加内容
            for key, value in content.items():
                doc.add_heading(key, level=1)
                doc.add_paragraph(str(value))
                
        return doc
        
    def _convert_to_html(self, content: Any, options: dict) -> str:
        """
        转换为HTML格式
        """
        if isinstance(content, str):
            # 如果输入是Markdown格式，转换为HTML
            if options.get('from_markdown', True):
                return markdown.markdown(content)
            else:
                return f"<html><body><pre>{content}</pre></body></html>"
        elif isinstance(content, dict):
            # 将字典转换为HTML表格
            html = "<table border='1'><tr><th>Key</th><th>Value</th></tr>"
            for key, value in content.items():
                html += f"<tr><td>{key}</td><td>{value}</td></tr>"
            html += "</table>"
            return html
        else:
            return f"<html><body><pre>{str(content)}</pre></body></html>"
            
    def _convert_to_json(self, content: Any, options: dict) -> dict:
        """
        转换为JSON格式
        """
        if isinstance(content, dict):
            return content
        else:
            return {"content": str(content)}
            
    def _convert_to_yaml(self, content: Any, options: dict) -> dict:
        """
        转换为YAML格式
        """
        return self._convert_to_json(content, options)
        
    def get_supported_formats(self) -> list:
        """
        获取支持的格式列表
        
        Returns:
            list: 支持的文件格式列表
        """
        return list(self.supported_formats.keys()) 